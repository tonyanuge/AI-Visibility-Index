"""Acceptance tests for the branded report renderer (build from the synthetic fixture)."""
import io, json, zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))
from docx import Document
from avix.report import ReportData, build_report, load_branding, builder as builder_mod

FIX = Path(__file__).resolve().parent / "fixtures" / "sample_report.json"
BRANDING = load_branding()


def _data():
    return ReportData.from_dict(json.loads(FIX.read_text(encoding="utf-8")))


def _docx_bytes():
    return build_report(_data(), BRANDING)


def _all_text(doc):
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _all_runs(doc):
    runs = list(p for para in doc.paragraphs for p in para.runs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    runs.extend(para.runs)
    return runs


def test_output_is_valid_docx():
    b = _docx_bytes()
    # opens via python-docx
    doc = Document(io.BytesIO(b))
    assert len(doc.paragraphs) > 0
    # unzips + parses as valid OOXML
    with zipfile.ZipFile(io.BytesIO(b)) as z:
        ET.fromstring(z.read("word/document.xml"))


def test_pagination_field_present():
    b = _docx_bytes()
    with zipfile.ZipFile(io.BytesIO(b)) as z:
        footer = next((n for n in z.namelist() if n.startswith("word/footer")), None)
        assert footer, "no footer part found"
        xml = z.read(footer).decode("utf-8")
    assert "PAGE" in xml and "NUMPAGES" in xml      # real fields, not static text
    assert "Page 1 of" not in xml                    # not hard-coded


def test_fields_marked_dirty_and_update_on_open():
    b = _docx_bytes()
    with zipfile.ZipFile(io.BytesIO(b)) as z:
        footer = next(n for n in z.namelist() if n.startswith("word/footer"))
        footer_xml = z.read(footer).decode("utf-8")
        settings_xml = z.read("word/settings.xml").decode("utf-8")
    # PAGE/NUMPAGES begin fldChar marked dirty so Word/Google Docs recompute on open
    assert 'w:dirty="true"' in footer_xml
    # document-wide update-fields-on-open flag
    assert "updateFields" in settings_xml


def test_verdict_colours_applied():
    doc = Document(io.BytesIO(_docx_bytes()))
    hexes = {str(r.font.color.rgb) for r in _all_runs(doc) if r.font.color and r.font.color.rgb}
    inv = BRANDING["verdict_colors"]["INVISIBLE"].upper()
    strong = BRANDING["verdict_colors"]["STRONG"].upper()
    assert inv in hexes, f"INVISIBLE colour {inv} not applied to any run"
    assert strong in hexes, f"STRONG colour {strong} not applied to any run"


def test_competitor_table_present():
    doc = Document(io.BytesIO(_docx_bytes()))
    found = False
    for t in doc.tables:
        headers = [c.text.strip() for c in t.rows[0].cells]
        if headers == ["Firm", "Times recommended ahead", "Assistants"]:
            assert len(t.rows) >= 2, "competitor table has no data rows"
            found = True
    assert found, "competitor table with expected headers not found"


def test_competitor_table_sorted_desc():
    doc = Document(io.BytesIO(_docx_bytes()))
    for t in doc.tables:
        headers = [c.text.strip() for c in t.rows[0].cells]
        if headers == ["Firm", "Times recommended ahead", "Assistants"]:
            counts = [int(r.cells[1].text) for r in t.rows[1:]]
            assert counts == sorted(counts, reverse=True)


def test_snapshot_date_on_cover():
    doc = Document(io.BytesIO(_docx_bytes()))
    assert _data().snapshot_date in _all_text(doc)


def test_disclaimer_present_no_overclaims():
    doc = Document(io.BytesIO(_docx_bytes()))
    text = _all_text(doc).lower()
    # protective wording REQUIRED
    assert "not a guarantee of future ranking" in text
    # affirmative overclaims BANNED
    for banned in ["guaranteed ranking", "we guarantee", "guaranteed results", "ai-powered"]:
        assert banned not in text, f"banned overclaim present: {banned}"


def test_headline_stat_matches_structured_data():
    data = _data()
    doc = Document(io.BytesIO(_docx_bytes()))
    text = _all_text(doc)
    # derived, so it always matches the tables (R7) — recommendation-based framing (no "checks")
    assert f"{data.mention_rate_pct}% share" in text
    assert f"of {data.total_checks} captured AI recommendations" in text
    assert "checks" not in text.lower() and "mention rate" not in text.lower()
    assert data.mention_rate_pct == 13 and data.total_checks == 15


def test_logo_optional():
    # the generator must run with no logo asset present
    assert not builder_mod.LOGO.exists(), "test expects no logo.png committed"
    Document(io.BytesIO(_docx_bytes()))  # builds without crashing


def test_fixture_is_synthetic():
    assert "SAMPLE" in _data().client
