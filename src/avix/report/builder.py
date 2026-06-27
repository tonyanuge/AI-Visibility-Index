"""Branded .docx renderer for AI Visibility Check reports.

build_report(data, branding) -> bytes. Config-driven: company, fonts, colours, footer and
disclaimer all come from config/report.yaml. The only "logic" here is layout + deriving
headline figures from the structured data (so they always match the tables)."""
import io
from pathlib import Path
import yaml

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .. import config
from .models import ReportData

ASSETS = config.CONFIG_DIR / "assets"
LOGO = ASSETS / "logo.png"


def load_branding() -> dict:
    """Load config/report.yaml. Kept separate so callers (CLI, tests) share one source."""
    with open(config.CONFIG_DIR / "report.yaml", "r", encoding="utf-8") as f:
        return yaml.safe_load(f) or {}


# ---------- low-level OOXML helpers (python-docx has no native API for these) ----------

def _add_field(paragraph, code: str):
    """Append a Word field (e.g. PAGE / NUMPAGES) as a real updating field, not static text.

    The begin fldChar is marked w:dirty="true" so Word / Google Docs recompute the field
    on open (without it they can show blank until manually refreshed). Pair with
    _enable_update_fields(), which sets the document-wide update-on-open flag."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin"); begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = f" {code} "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(begin); run._r.append(instr); run._r.append(end)


def _enable_update_fields(doc):
    """Add <w:updateFields w:val="true"/> to settings so Word/Google Docs refresh fields on open."""
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        el = OxmlElement("w:updateFields")
        el.set(qn("w:val"), "true")
        settings.append(el)


def _shade(cell, fill_hex: str):
    """Apply solid cell background shading via w:shd (clear pattern, explicit fill)."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_margins(table, top=80, bottom=80, left=110, right=110):
    """Table-wide cell padding (twentieths of a point) for breathing room in tables."""
    tblPr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)


# ---------- formatting helpers ----------

def _font(run, branding, size=11, bold=False, color=None, italic=False):
    run.font.name = branding.get("font_family", "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def _para(doc, branding, text="", size=11, bold=False, color=None, italic=False,
          space_after=6, space_before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        p.alignment = align
    if text:
        _font(p.add_run(text), branding, size, bold, color, italic)
    return p


def _heading(doc, branding, text, size=14):
    colors = branding.get("colors", {})
    return _para(doc, branding, text, size=size, bold=True,
                 color=colors.get("cover_rule", "0E7C7B"), space_before=12, space_after=4)


def _bullets(doc, branding, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        _font(p.add_run(it), branding, 11)


def _verdict_color(branding, state: str) -> str:
    return branding.get("verdict_colors", {}).get(state, branding.get("colors", {}).get("ink", "142128"))


def _verdict_label(branding, state: str) -> str:
    return branding.get("verdict_labels", {}).get(state, state.replace("_", " ").title())


# ---------- the report ----------

def build_report(data: ReportData, branding: dict) -> bytes:
    doc = Document()
    colors = branding.get("colors", {})
    company = branding.get("company", "")

    # base font for the Normal style
    normal = doc.styles["Normal"]
    normal.font.name = branding.get("font_family", "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(colors.get("ink", "142128"))

    _build_footer(doc, branding, company, data.client)
    _enable_update_fields(doc)  # refresh PAGE/NUMPAGES fields on open (Word / Google Docs)

    # ---- unmissable SAMPLE strip at the very top (demo reports only) ----
    if getattr(data, "sample_banner", ""):
        _sample_strip(doc, branding, data.sample_banner)

    # ---- R1 optional logo (cover) ----
    if LOGO.exists():
        try:
            doc.add_picture(str(LOGO), width=Inches(1.6))
        except Exception:
            pass  # never let a bad asset crash a report

    # ---- R2 cover / title block ----
    _para(doc, branding, f"AI Visibility Check — {data.client}", size=22, bold=True,
          color=colors.get("ink", "142128"), space_after=2)
    _kv(doc, branding, "Market", data.market)
    _kv(doc, branding, "Area", data.area)
    _kv(doc, branding, "Snapshot date", data.snapshot_date, value_size=13, value_bold=True)  # prominent
    _kv(doc, branding, "Prepared by", data.prepared_by)
    # overall verdict, coloured (R3)
    pv = _para(doc, branding, "Overall verdict: ", bold=True, space_after=4)
    _font(pv.add_run(data.overall_verdict_label), branding, 11, bold=True,
          color=_verdict_color(branding, data.overall_verdict_state))

    # ---- R4 headline-stat callout (computed from structured data, R7) ----
    _callout(doc, branding,
             f"{data.mentioned_count} of {data.total_checks} captured AI recommendations "
             f"name this business ({data.mention_rate_pct}% share)")

    # ---- 1. Executive Summary ----
    _heading(doc, branding, "1. Executive Summary")
    for para in data.executive_summary:
        _para(doc, branding, para)

    # ---- 2. What We Tested ----
    _heading(doc, branding, "2. What We Tested")
    _para(doc, branding, "Search prompts tested:", bold=True, space_after=2)
    _bullets(doc, branding, data.what_we_tested_prompts)
    _para(doc, branding, "AI assistants tested:", bold=True, space_after=2)
    _bullets(doc, branding, data.what_we_tested_assistants)
    _para(doc, branding, f"How the captured recommendations break down ({data.total_checks} in total):",
          bold=True, space_after=4)
    _sample_table(doc, branding, data)

    # ---- 3. Key Findings ----
    _heading(doc, branding, "3. Key Findings")
    _bullets(doc, branding, data.key_findings)
    if data.competitor_counts:
        _para(doc, branding, "Competitors recommended ahead:", bold=True,
              space_before=6, space_after=4)
        _competitor_table(doc, branding, data)

    # ---- 4. Visibility Verdict (table + key + supporting line) ----
    _heading(doc, branding, "4. Visibility Verdict")
    _verdict_table(doc, branding, data)
    _verdict_key(doc, branding)
    _para(doc, branding,
          f"Supporting numbers from this sample: {data.total_checks} captured AI recommendations · "
          f"named {data.mentioned_count} · invisible {data.invisible_count} · "
          f"visible-but-beaten {data.beaten_count} · share of recommendations ≈ {data.mention_rate_pct}%.",
          size=10, color=colors.get("muted", "5D6F76"), space_before=8, space_after=10)

    # ---- 5. What This Means ----
    _heading(doc, branding, "5. What This Means")
    for para in data.what_this_means:
        _para(doc, branding, para)

    # ---- 6. Scope of this report (renders data.recommended_actions as scope notes) ----
    _heading(doc, branding, "6. Scope of this report")
    _bullets(doc, branding, data.recommended_actions)

    # ---- 7. Disclaimer / Boundary (R6 verbatim) ----
    _heading(doc, branding, "7. Disclaimer / Boundary")
    _para(doc, branding, branding.get("disclaimer", "").strip(), size=10,
          color=colors.get("muted", "5D6F76"))
    _para(doc, branding, branding.get("confidential_note", "").strip(), size=9,
          color=colors.get("muted", "5D6F76"), italic=True, space_before=4)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------- section builders ----------

def _build_footer(doc, branding, company, client):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = branding.get("footer_confidential", "{company} · Confidential — prepared for {client}")
    left = left.format(company=company, client=client)
    _font(p.add_run(left), branding, 8, color=branding.get("colors", {}).get("muted", "5D6F76"))
    p.add_run("\t")
    # right: real "Page X of Y" field
    _font(p.add_run("Page "), branding, 8, color=branding.get("colors", {}).get("muted", "5D6F76"))
    _add_field(p, "PAGE")
    _font(p.add_run(" of "), branding, 8, color=branding.get("colors", {}).get("muted", "5D6F76"))
    _add_field(p, "NUMPAGES")


def _kv(doc, branding, label, value, value_size=11, value_bold=False):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    _font(p.add_run(f"{label}: "), branding, 11, bold=True)
    _font(p.add_run(value), branding, value_size, bold=value_bold)


def _sample_strip(doc, branding, text):
    """Full-width red strip with white bold text — marks the whole report as sample data."""
    red = branding.get("verdict_colors", {}).get("INVISIBLE", "C0392B")
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_cell_margins(table, top=110, bottom=110, left=160, right=160)
    cell = table.rows[0].cells[0]
    _shade(cell, red)
    cell.width = Inches(6.6)
    _font(cell.paragraphs[0].add_run(text), branding, 12, bold=True, color="FFFFFF")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _callout(doc, branding, text):
    colors = branding.get("colors", {})
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_cell_margins(table, top=120, bottom=120, left=160, right=160)
    cell = table.rows[0].cells[0]
    _shade(cell, colors.get("callout_bg", "EEF4F3"))
    cell.width = Inches(6.4)
    para = cell.paragraphs[0]
    _font(para.add_run(text), branding, 15, bold=True, color=colors.get("cover_rule", "0E7C7B"))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _sample_table(doc, branding, data):
    colors = branding.get("colors", {})
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_cell_margins(table)
    widths = [Inches(1.3), Inches(2.7), Inches(2.0), Inches(0.7)]
    headers = ["Search theme", "Prompt", "Assistants captured", "Count"]
    for c, txt, w in zip(table.rows[0].cells, headers, widths):
        c.width = w
        _shade(c, colors.get("callout_bg", "EEF4F3"))
        _font(c.paragraphs[0].add_run(txt), branding, 10, bold=True)
    for r in data.sample_rows:
        cells = table.add_row().cells
        vals = [r.theme, r.prompt, r.assistants, str(r.checks)]
        for c, txt, w in zip(cells, vals, widths):
            c.width = w
            _font(c.paragraphs[0].add_run(txt), branding, 10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _competitor_table(doc, branding, data):
    """R5: competitors as a 3-column table, sorted by frequency descending."""
    colors = branding.get("colors", {})
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_cell_margins(table)
    widths = [Inches(2.9), Inches(1.7), Inches(2.0)]
    headers = ["Firm", "Times recommended ahead", "Assistants"]
    for c, txt, w in zip(table.rows[0].cells, headers, widths):
        c.width = w
        _shade(c, colors.get("callout_bg", "EEF4F3"))
        _font(c.paragraphs[0].add_run(txt), branding, 10, bold=True)
    for comp in sorted(data.competitor_counts, key=lambda x: x.times_ahead, reverse=True):
        cells = table.add_row().cells
        vals = [comp.firm, str(comp.times_ahead), comp.assistants]
        for c, txt, w in zip(cells, vals, widths):
            c.width = w
            _font(c.paragraphs[0].add_run(txt), branding, 10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _verdict_table(doc, branding, data):
    colors = branding.get("colors", {})
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_cell_margins(table)
    widths = [Inches(4.6), Inches(2.1)]
    for c, txt, w in zip(table.rows[0].cells, ["Search type", "Verdict"], widths):
        c.width = w
        _shade(c, colors.get("callout_bg", "EEF4F3"))
        _font(c.paragraphs[0].add_run(txt), branding, 10, bold=True)
    for label, state in data.verdicts.items():
        cells = table.add_row().cells
        cells[0].width = widths[0]; cells[1].width = widths[1]
        _font(cells[0].paragraphs[0].add_run(label), branding, 10)
        _font(cells[1].paragraphs[0].add_run(_verdict_label(branding, state)),
              branding, 10, bold=True, color=_verdict_color(branding, state))
    # overall row, coloured
    cells = table.add_row().cells
    cells[0].width = widths[0]; cells[1].width = widths[1]
    _font(cells[0].paragraphs[0].add_run("Overall AI visibility"), branding, 10, bold=True)
    _font(cells[1].paragraphs[0].add_run(data.overall_verdict_label), branding, 10, bold=True,
          color=_verdict_color(branding, data.overall_verdict_state))


def _verdict_key(doc, branding):
    """Compact legend so the verdict colour system is documented (and every colour appears)."""
    states = ["INVISIBLE", "VISIBLE_BEATEN", "STRONG"]
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(0)
    _font(p.add_run("Verdict key: "), branding, 9, bold=True,
          color=branding.get("colors", {}).get("muted", "5D6F76"))
    for i, st in enumerate(states):
        if i:
            _font(p.add_run("   ·   "), branding, 9, color=branding.get("colors", {}).get("muted", "5D6F76"))
        _font(p.add_run(_verdict_label(branding, st)), branding, 9, bold=True,
              color=_verdict_color(branding, st))
